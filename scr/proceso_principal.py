from __future__ import annotations

from dataclasses import dataclass
from enum import Enum
import json
from pathlib import Path
import re

from motor_ocr import MotorOCRPaddle
from procesador_imagen import ProcesadorImagenOpenCV
from procesador_pdf import ProcesadorPDF


class TipoArchivo(str, Enum):
    PDF = "pdf"
    PNG = "png"
    JPG = "jpg"
    JPEG = "jpeg"


EXTENSIONES_PERMITIDAS = {item.value for item in TipoArchivo}

# Flujo actual: procesar un solo archivo definido por ruta directa.
RUTA_ARCHIVO_ENTRADA = "entrada/ejemplo.pdf"  # Cambia esta ruta para probar con diferentes archivos.


@dataclass(frozen=True)
class ArchivoEntrada:
    nombre: str
    ruta: Path
    tipo: TipoArchivo


def detectar_tipo_archivo(file_path: Path) -> TipoArchivo:
    extension = file_path.suffix.lower().lstrip(".")
    if extension not in EXTENSIONES_PERMITIDAS:
        raise ValueError(
            f"Tipo no soportado para {file_path.name}. "
            f"Permitidos: {', '.join(sorted(EXTENSIONES_PERMITIDAS))}"
        )
    return TipoArchivo(extension)


def cargar_archivo_directo(raiz_proyecto: Path) -> ArchivoEntrada:
    ruta_archivo = (raiz_proyecto / RUTA_ARCHIVO_ENTRADA).resolve()
    if not ruta_archivo.exists() or not ruta_archivo.is_file():
        raise FileNotFoundError(f"No existe el archivo configurado: {ruta_archivo}")

    tipo = detectar_tipo_archivo(ruta_archivo)
    return ArchivoEntrada(nombre=ruta_archivo.name, ruta=ruta_archivo, tipo=tipo)


class EnrutadorProcesamiento:
    def __init__(self) -> None:
        self.motor_ocr = MotorOCRPaddle()
        self.procesador_pdf = ProcesadorPDF()
        self.procesador_imagen = ProcesadorImagenOpenCV(motor_ocr=self.motor_ocr)

    def procesar(self, archivo: ArchivoEntrada, carpeta_salida: Path) -> dict:
        if archivo.tipo == TipoArchivo.PDF:
            resultado_pdf = self.procesador_pdf.procesar(archivo.ruta, carpeta_salida)
            rutas_png = resultado_pdf.get("rutas_png_para_ocr")
            if not rutas_png and resultado_pdf.get("ruta_png_para_ocr"):
                rutas_png = [resultado_pdf.get("ruta_png_para_ocr")]

            resultados_imagenes = []
            for ruta_png in rutas_png or []:
                ruta_png_path = Path(ruta_png)
                resultados_imagenes.append(self.procesador_imagen.procesar(ruta_png_path, carpeta_salida))

            resultado_pdf["ocr"] = [resultado.get("ocr", {}) for resultado in resultados_imagenes]
            resultado_pdf["nota"] = f'{resultado_pdf.get("nota", "")} -> OCR aplicado con PaddleOCR en todas las paginas'
            resultado_pdf["paginas_procesadas_ocr"] = len(resultados_imagenes)

            salida_texto = resultado_pdf.get("salida_texto")
            if salida_texto:
                texto_final = _extraer_markdown_o_texto(resultado_pdf)
                Path(str(salida_texto)).write_text(_limpiar_html_de_tabla_para_txt(texto_final), encoding="utf-8")
            return resultado_pdf

        if archivo.tipo in {TipoArchivo.PNG, TipoArchivo.JPG, TipoArchivo.JPEG}:
            return self.procesador_imagen.procesar(archivo.ruta, carpeta_salida)

        raise ValueError(f"No existe procesador para el tipo: {archivo.tipo.value}")


def guardar_resultado_json(resultado: dict, carpeta_salida: Path, nombre_base: str) -> Path:
    ruta_salida = carpeta_salida / f"resultado_{nombre_base}.json"
    ruta_salida.write_text(json.dumps(_serializar(resultado), ensure_ascii=False, indent=2), encoding="utf-8")
    return ruta_salida


def guardar_resultado_word(resultado: dict, carpeta_salida: Path, nombre_base: str) -> Path:
    from docx import Document  # type: ignore

    texto_plano = _extraer_markdown_o_texto(resultado)
    estructura = _serializar(_extraer_estructura(resultado))
    ruta_salida = carpeta_salida / f"ocr_{nombre_base}.docx"

    document = Document()
    document.add_heading("Resultado OCR", level=1)
    document.add_paragraph(f"Archivo: {resultado.get('archivo', 'N/A')}")
    document.add_paragraph(f"Motor: {_extraer_motor(resultado)}")

    document.add_heading("Texto OCR / Markdown", level=2)
    _agregar_bloques_al_documento(document, estructura, texto_plano)

    rutas_paginas = resultado.get("salida_png_500dpi")
    if isinstance(rutas_paginas, list) and rutas_paginas:
        document.add_heading("Vista Original", level=2)
        for ruta in rutas_paginas:
            try:
                document.add_picture(str(ruta))
            except Exception:
                continue
    try:
        document.save(str(ruta_salida))
        return ruta_salida
    except PermissionError:
        ruta_alternativa = carpeta_salida / f"ocr_{nombre_base}_nuevo.docx"
        document.save(str(ruta_alternativa))
        return ruta_alternativa


def _agregar_bloques_al_documento(document, estructura, texto_plano: str) -> None:
    if isinstance(estructura, list) and estructura:
        for bloque in estructura:
            if not isinstance(bloque, dict):
                continue
            texto = str(bloque.get("texto", "")).strip()
            if not texto:
                continue
            tipo = str(bloque.get("tipo", "paragraph"))
            etiqueta = str(bloque.get("label_layout", "")).lower()

            if tipo == "table" or "table" in etiqueta or "<table" in texto.lower():
                if _agregar_tabla_html_si_existe(document, texto):
                    continue

            if tipo == "heading":
                document.add_heading(texto, level=2)
            else:
                document.add_paragraph(texto)
        return

    if texto_plano:
        for linea in texto_plano.splitlines():
            linea = linea.strip()
            if not linea:
                continue
            if linea.startswith("## "):
                document.add_heading(linea[3:].strip(), level=2)
            else:
                document.add_paragraph(linea)
        return

    document.add_paragraph("(sin texto detectado)")


def _extraer_markdown_o_texto(resultado: dict) -> str:
    ocr = resultado.get("ocr")
    if isinstance(ocr, dict) and ocr.get("markdown_plano"):
        return str(ocr.get("markdown_plano"))

    if isinstance(ocr, list):
        partes = []
        for item in ocr:
            if isinstance(item, dict):
                markdown = item.get("markdown_plano")
                texto = item.get("texto_plano")
                if markdown:
                    partes.append(str(markdown))
                elif texto:
                    partes.append(str(texto))
        if partes:
            return "\n\n".join(partes)

    ocr = resultado.get("ocr")
    if isinstance(ocr, dict) and ocr.get("texto_plano"):
        texto = ocr.get("texto_plano")
        if isinstance(texto, list):
            return "\n".join(str(item) for item in texto if str(item).strip())
        return str(texto)

    if resultado.get("texto_plano"):
        texto = resultado.get("texto_plano")
        if isinstance(texto, list):
            return "\n".join(str(item) for item in texto if str(item).strip())
        return str(texto)

    return ""


def _extraer_motor(resultado: dict) -> str:
    ocr = resultado.get("ocr")
    if isinstance(ocr, dict) and ocr.get("motor"):
        return str(ocr.get("motor"))

    if isinstance(ocr, list):
        motores = []
        for item in ocr:
            if isinstance(item, dict) and item.get("motor"):
                motores.append(str(item.get("motor")))
        if motores:
            return ", ".join(motores)

    return str(resultado.get("motor", "N/A"))


def _extraer_estructura(resultado: dict):
    ocr = resultado.get("ocr")
    if isinstance(ocr, dict) and ocr.get("estructura_documento") is not None:
        return ocr.get("estructura_documento")

    if isinstance(ocr, list):
        estructuras = []
        for item in ocr:
            if isinstance(item, dict) and item.get("estructura_documento"):
                estructuras.extend(item.get("estructura_documento") or [])
        if estructuras:
            return estructuras

    return []


def _agregar_tabla_html_si_existe(document, texto: str) -> bool:
    if "<table" not in texto.lower():
        return False

    filas = re.findall(r"<tr>(.*?)</tr>", texto, flags=re.IGNORECASE | re.DOTALL)
    if not filas:
        return False

    tabla_parseada: list[list[str]] = []
    for fila in filas:
        celdas = re.findall(r"<t[dh]>(.*?)</t[dh]>", fila, flags=re.IGNORECASE | re.DOTALL)
        celdas_limpias = [re.sub(r"<[^>]+>", "", celda).strip() for celda in celdas]
        if celdas_limpias:
            tabla_parseada.append(celdas_limpias)

    if not tabla_parseada:
        return False

    total_columnas = max(len(fila) for fila in tabla_parseada)
    tabla = document.add_table(rows=len(tabla_parseada), cols=total_columnas)
    tabla.style = "Table Grid"

    for i, fila in enumerate(tabla_parseada):
        for j, valor in enumerate(fila):
            tabla.cell(i, j).text = valor

    return True


def _limpiar_html_de_tabla_para_txt(texto: str) -> str:
    if "<table" not in texto.lower():
        return texto

    def reemplazo_tabla(match):
        tabla_html = match.group(0)
        filas = re.findall(r"<tr>(.*?)</tr>", tabla_html, flags=re.IGNORECASE | re.DOTALL)
        lineas = []
        for fila in filas:
            celdas = re.findall(r"<t[dh]>(.*?)</t[dh]>", fila, flags=re.IGNORECASE | re.DOTALL)
            celdas_limpias = [re.sub(r"<[^>]+>", "", celda).strip() for celda in celdas]
            if celdas_limpias:
                lineas.append(" | ".join(celdas_limpias))
        return "\n".join(lineas)

    return re.sub(
        r"<table.*?>.*?</table>",
        reemplazo_tabla,
        texto,
        flags=re.IGNORECASE | re.DOTALL,
    )


def _serializar(valor):
    if isinstance(valor, Path):
        return str(valor)
    if isinstance(valor, dict):
        return {str(clave): _serializar(v) for clave, v in valor.items()}
    if isinstance(valor, list):
        return [_serializar(item) for item in valor]
    if isinstance(valor, tuple):
        return [_serializar(item) for item in valor]
    return valor


def main() -> None:
    raiz_proyecto = Path(__file__).resolve().parent.parent
    carpeta_salida = raiz_proyecto / "salida"
    carpeta_salida.mkdir(parents=True, exist_ok=True)

    archivo = cargar_archivo_directo(raiz_proyecto)
    enrutador = EnrutadorProcesamiento()
    resultado = enrutador.procesar(archivo, carpeta_salida)
    json_generado = guardar_resultado_json(resultado, carpeta_salida, archivo.ruta.stem)
    docx_generado = guardar_resultado_word(resultado, carpeta_salida, archivo.ruta.stem)

    print("Archivo configurado detectado correctamente:")
    print(f"- {archivo.nombre} -> tipo={archivo.tipo.value}")
    print(f"Resultado guardado en: {json_generado}")
    print(f"Documento Word guardado en: {docx_generado}")
    for clave in ("salida_pdf", "salida_texto", "salida_preview"):
        valor = resultado.get(clave)
        if valor:
            print(f"- {clave}: {valor}")


if __name__ == "__main__":
    main()
